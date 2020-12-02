from docx import Document
import os
def result(src_file):
    document = Document()
    document.add_heading('                                                              Result')
    document.add_heading('  Grade:', level=2)
    paragraph = document.add_paragraph('            C')
    document.add_heading('Sentiment Analysis:')
    document.add_picture(src_file+'_sentiment.png')
    document.add_heading('Programming Languages:')
    f=open(src_file+"_pro_lang.txt",'r')
    temp=f.read()
    f.close()
    l=temp.split("\n")
    del(l[-1])
    for x in l:
        paragraph = document.add_paragraph(x,style='ListBullet')
    document.add_page_break()
    document.add_heading('Spelling Mistakes:')
    document.add_heading('Incorrect                                                 Correct')
    f=open(src_file+"_spell_ext.txt")
    temp=f.read()
    l=temp.split("\n")
    del(l[-1])
    table = document.add_table(rows=len(l), cols=2)
    i=int(0)
    for x in l:
        er=x.split(":")
        for j in range(0,2):
            cell = table.cell(i,j)
            cell.text=er[j]
        i+=1
    document.add_heading('Subject Of Intrest:')
    f=open(src_file+"_sub.txt","r")
    temp=f.read()
    l=temp.split("\n")
    del(l[-1])
    for x in l:
        paragraph = document.add_paragraph(x,style='ListBullet')
    document.save("E:/xampp/htdocs/Parser/Output/"+src_file+".docx")
    '''
    os.remove(src_file+"_spell_ext.txt")
    os.remove(src_file+"_sent_ext.xml")
    os.remove(src_file+"_grade.txt")
    os.remove(src_file+"_pro_lang.txt")
    os.remove(src_file+'_sentiment.png')
    os.remove(src_file+'_sub.txt')'''

